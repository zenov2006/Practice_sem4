from pathlib import Path

import pdfplumber
from docx import Document
from fastapi import HTTPException, status


def extract_text_from_pdf(file_path: Path) -> list[dict[str, int | str]]:
    """
    Извлекает текст из PDF-документа постранично.

    Args:
        file_path: Путь к PDF-файлу.

    Returns:
        list[dict[str, int | str]]: Список страниц с номером страницы и текстом.

    Raises:
        HTTPException: Если PDF не удалось прочитать или текст не был найден.
    """
    pages_text: list[dict[str, int | str]] = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""

                if text.strip():
                    pages_text.append(
                        {
                            "page_number": page_index,
                            "text": text,
                        }
                    )
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не удалось прочитать PDF-файл",
        ) from exc

    if not pages_text:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не удалось извлечь текст из PDF-файла",
        )

    return pages_text


def extract_text_from_docx(file_path: Path) -> list[dict[str, int | str]]:
    """
    Извлекает текст из DOCX-документа.

    У DOCX нет стабильной информации о страницах, поэтому весь текст
    считается первой страницей. Для полнотекстового поиска этого достаточно.

    Args:
        file_path: Путь к DOCX-файлу.

    Returns:
        list[dict[str, int | str]]: Список с одним элементом: номер страницы и текст.

    Raises:
        HTTPException: Если DOCX не удалось прочитать или текст не был найден.
    """
    try:
        document = Document(file_path)

        text_parts: list[str] = []

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                text_parts.append(paragraph_text)

        for table in document.tables:
            for row in table.rows:
                row_cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if row_cells:
                    text_parts.append(" | ".join(row_cells))

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не удалось прочитать DOCX-файл",
        ) from exc

    text = "\n".join(text_parts)

    if not text.strip():
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не удалось извлечь текст из DOCX-файла",
        )

    return [
        {
            "page_number": 1,
            "text": text,
        }
    ]


def extract_text_from_document(file_path: Path) -> list[dict[str, int | str]]:
    """
    Определяет тип документа по расширению и извлекает из него текст.

    Args:
        file_path: Путь к загруженному документу.

    Returns:
        list[dict[str, int | str]]: Список страниц с текстом.

    Raises:
        HTTPException: Если формат файла не поддерживается.
    """
    file_extension = file_path.suffix.lower()

    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)

    if file_extension == ".docx":
        return extract_text_from_docx(file_path)

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Неподдерживаемый формат документа",
    )